"""
# Name of program: Automated Personalized Christmas Card Maker
# Version: python3
# Author: Joshua Sam
# Purpose: To automate the creation of personalized Christmas Cards
# Inputs: A document titled "names.txt" with each line containing a name
# Outputs: Creates a word document called "christmasCard.docx" with each page containing a personalized
           Christmas card with the names of the people in "names.txt"
"""

# using os and python-docx frameworks
import os
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.shared import Pt


def createCard(txtFile, docName):
    doc = docx.Document()

    intro = 'Merry Christmas!'
    line2 = 'Wishing you peace, joy, and all the best this wonderful holiday has to offer.'
    line3 = 'Best wishes,'
    conclusion = "Joshua"

    with open(txtFile) as guestList:
        for guest in guestList:
            name = guest[:-1]
            p1 = doc.add_paragraph()
            p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f1 = p1.add_run(intro)
            f1.font.bold = True
            f1.font.italic = True
            f1.font.size = Pt(20)

            p2 = doc.add_paragraph()
            p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f2 = p2.add_run(name)
            f2.font.size = Pt(12)

            p3 = doc.add_paragraph()
            p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f3 = p3.add_run(line2)
            f3.font.size = Pt(12)

            p4 = doc.add_paragraph()
            p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f4 = p4.add_run(line3)
            f4.font.size = Pt(12)

            p5 = doc.add_paragraph()
            p5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f5 = p5.add_run(conclusion)
            f5.font.size = Pt(12)

            doc.add_page_break()

    doc.save(docName)

if __name__ == "__main__":
    createCard('names.txt', 'christmasCard.docx')